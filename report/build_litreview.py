"""Two-page literature review for Prof. Ye. Monochrome, Times New Roman.

Design brief: he should be able to read it in five minutes and know what the
review decided. Everything that is not a decision, a number, or a citation has
been cut. All text is black; links are black and underlined.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

OUT = Path(__file__).parent / "BRSET_to_mBRSET_Literature_Review.docx"
FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(10.5)
st.font.color.rgb = BLACK
st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.8)
    s.top_margin = s.bottom_margin = Inches(0.65)


def fmt(run, size=10.5, bold=False, italic=False):
    run.font.name = FONT; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    run.font.color.rgb = BLACK


def P(text, size=10.5, bold=False, italic=False, after=5, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    fmt(p.add_run(text), size, bold, italic)
    return p


def H(text, size=11.5, before=9):
    return P(text, size, bold=True, after=3, before=before)


def link(par, url, text=None):
    r_id = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "000000"); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    f = OxmlElement("w:rFonts"); f.set(qn("w:ascii"), FONT); f.set(qn("w:hAnsi"), FONT); rPr.append(f)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "17"); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text or url; r.append(t)
    h.append(r); par._p.append(h)


def table(rows, widths, size=9.5):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.cell(i, j); cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            if j and i:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(v)) < 12 else WD_ALIGN_PARAGRAPH.LEFT
            fmt(p.add_run(str(v)), size, bold=(i == 0))
            cell.width = Inches(widths[j])
    for j, w in enumerate(widths):
        for row in t.rows:
            row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


# ── title ────────────────────────────────────────────────────────────────
P("Improving BRSET to mBRSET Transfer: Literature Review and Proposed Direction",
  14, bold=True, after=2)
P("Sahith Reddy Thummala.  Prepared for Prof. Dong Hye Ye.  21 August 2026.", 9.5, after=9)

# ── 1 ────────────────────────────────────────────────────────────────────
H("1. The problem", before=0)
P("The ConvNeXt V2 baseline reaches AUC 0.9906 and class-macro F1 0.9374 on BRSET, which is tabletop-camera "
  "data from a hospital eye clinic. Applied unchanged to mBRSET, handheld phone-camera data from a community "
  "diabetes screening campaign, it misses 50 of 159 diseased patients instead of 10 of 162. No mBRSET image "
  "is used in training.")
table([
    ["Diabetic retinopathy", "On BRSET", "On mBRSET"],
    ["AUC (ranking ability)", "0.9906", "0.9060"],
    ["F1, class-macro", "0.9374", "0.8668"],
    ["Recall, diseased class", "0.9383", "0.6855"],
    ["Diseased cases missed", "10 of 162", "50 of 159"],
], [2.4, 2.3, 2.3])
P("Ranking barely suffers; the decision collapses. Two things change at once. The camera degrades, with 83 "
  "percent of mBRSET images carrying an artifact, and the population changes, from 6.6 percent with DR to "
  "23.3 percent. Most published cross-device work changes only one of these.")

# ── 2 ────────────────────────────────────────────────────────────────────
H("2. Where the gap lives, measured rather than assumed")
table([
    ["Diabetic retinopathy, diseased-class F1 on mBRSET", "F1"],
    ["Transferred as is", "0.7842"],
    ["Best reachable at any decision threshold", "0.7927"],
    ["Fine-tuned on labelled mBRSET", "0.8317"],
], [4.7, 2.3])
P("Sweeping every threshold from 0.005 to 0.995 shows the best any threshold can reach is 0.7927. "
  "Calibration, temperature scaling and label-shift correction only rescale scores, so each selects a point "
  "on that same curve and none can exceed it. That family is ruled out by measurement rather than by trial. "
  "Of the remaining gap, 18 percent is the decision threshold and 82 percent is the representation.")
P("The shift is not only a change in prevalence. ROC is invariant to class balance [11], so if only the "
  "disease rate had moved, AUC would have held. It fell from 0.9906 to 0.9060, which means the images "
  "themselves differ.")

# ── 3 ────────────────────────────────────────────────────────────────────
H("3. What the review decides")
P("Peer-reviewed versions only, never preprints. Citation counts retrieved from the Semantic Scholar API by "
  "DOI on 21 August 2026; four counts could not be retrieved because of rate limiting and are marked so "
  "rather than estimated. Six themes were searched independently.", 9.5)
table([
    ["Approach", "Evidence", "Verdict for this project"],
    ["Degradation and appearance\naugmentation", "Tellez, Med. Image Anal. 2019, 665 cites [1];\nChe, MICCAI 2023, 46 cites [2]",
     "Adopt. The only intervention with peer-reviewed\nevidence of raising AUC under real fundus device shift"],
    ["Adversarial adaptation\n(DANN, MDD, CycleGAN)", "Lin, MICCAI 2022, 5 cites [3]",
     "Eliminated. Documented to collapse on real\ntabletop to portable fundus data"],
    ["Domain generalization\nmethods", "Gulrajani and Lopez-Paz, ICLR 2021,\n1,504 cites [4]",
     "Treat with caution. None reliably beats a\nwell-tuned plain baseline"],
    ["Test-time adaptation", "Boudiaf, CVPR 2022, 257 cites [5]",
     "Eliminated. Loses up to 66 points under prior\nshift, which is the regime here"],
    ["Threshold and calibration\ncorrection", "Roschewitz, Nature Comms 2023, 43 cites [6];\nBBSE, ICML 2018, 679 cites [8]",
     "Capped. Cheap, but provably limited to F1 0.7927\nhere, and BBSE assumes a single-label softmax"],
], [1.5, 2.6, 2.9])
P("The value of the review was negative. Three of the four approaches I would have reached for first are "
  "published failures in this exact setting, which removed them before any implementation effort.")

# ── 4 ────────────────────────────────────────────────────────────────────
H("4. Proposed direction")
P("GDRNet's FundusAug [2] applies nine operations, five colour transformations and four degradations, each "
  "at a fixed probability of 0.5 with hand-chosen magnitudes. It is generic damage for generic robustness, "
  "calibrated to no particular camera. However, 4,272 of the 5,164 mBRSET images already carry artifact "
  "annotations, so the handheld device's actual degradation statistics are measurable without using a single "
  "diagnostic label.")
P("I propose fitting the augmentation to those measured statistics, using the physically grounded fundus "
  "degradation model of Shen et al. [7], which derives light transmission disturbance, blur and retinal "
  "artifacts from the optics rather than by hand. The required control is generic FundusAug at its published "
  "settings; if fitting to the target device does not beat guessing, the contribution fails and I will report "
  "that. The gap decomposition in Section 2 supplies the budget for each intervention before either is built.")
P("Closest prior work, stated plainly. Joseph et al. [10] adapt images across cameras for portable DR "
  "screening, but in the opposite direction, translating target images toward the source, and never treat the "
  "decision cutoff. Roschewitz et al. [9] identify which kind of shift has occurred but do not quantify what "
  "each component costs, nor establish a ceiling. Park et al. [12] handle covariate and label shift jointly "
  "but only for single-label softmax outputs, whereas DR and macular edema co-occur here.")

# ── references ───────────────────────────────────────────────────────────
H("References")
refs = [
 ("Tellez et al. Quantifying the effects of data augmentation and stain color normalization in CNNs for computational pathology. Medical Image Analysis 58:101544, 2019. 665 cites.", "https://doi.org/10.1016/j.media.2019.101544"),
 ("Che, Cheng, Jin, Chen. Towards Generalizable Diabetic Retinopathy Grading in Unseen Domains (GDRNet). MICCAI 2023, 430-440. 46 cites.", "https://doi.org/10.1007/978-3-031-43904-9_42"),
 ("Lin, Shi, Zhang, Shang, He, Ge. Camera Adaptation for Fundus-Image-Based CVD Risk Estimation. MICCAI 2022, 593-603. 5 cites.", "https://doi.org/10.1007/978-3-031-16434-7_57"),
 ("Gulrajani, Lopez-Paz. In Search of Lost Domain Generalization. ICLR 2021. 1,504 cites.", "https://openreview.net/forum?id=lQdXeXDoWtI"),
 ("Boudiaf, Mueller, Ben Ayed, Bertinetto. Parameter-free Online Test-time Adaptation (LAME). CVPR 2022, 8344-8353. 257 cites.", "https://doi.org/10.1109/CVPR52688.2022.00816"),
 ("Roschewitz, Khara, Yearsley et al. Automatic correction of performance drift under acquisition shift in medical image classification. Nature Communications 14:6608, 2023. 43 cites.", "https://doi.org/10.1038/s41467-023-42396-y"),
 ("Shen, Fu, Shen, Shao. Modeling and Enhancing Low-Quality Retinal Fundus Images (cofe-Net). IEEE Transactions on Medical Imaging 40(3), 2021. Count not retrieved.", "https://arxiv.org/abs/2005.05594"),
 ("Lipton, Wang, Smola. Detecting and Correcting for Label Shift with Black Box Predictors (BBSE). ICML 2018. 679 cites.", "https://proceedings.mlr.press/v80/lipton18a.html"),
 ("Roschewitz et al. Automatic Dataset Shift Identification to Support Safe Deployment of Medical Imaging AI. MICCAI 2025. Count not retrieved.", "https://doi.org/10.1007/978-3-032-04981-0_7"),
 ("Joseph, Chen, Liu, Zhu, Ramasamy, Ravilla, Ge, He. Enhancing AI-based diabetic retinopathy diagnosis through universal cross-camera image adaptation. BMJ Open Ophthalmology 10(1), 2025. Count not retrieved.", "https://doi.org/10.1136/bmjophth-2025-002238"),
 ("Fawcett. An introduction to ROC analysis. Pattern Recognition Letters 27(8):861-874, 2006. 21,726 cites.", "https://doi.org/10.1016/j.patrec.2005.10.010"),
 ("Park, Yang, Choo, Yun. Label Shift Adapter for Test-Time Adaptation under Covariate and Label Shifts. ICCV 2023, 16421-16431. Count not retrieved.", "https://arxiv.org/abs/2308.08810"),
 ("Nakayama, Restrepo et al. BRSET: A Brazilian Multilabel Ophthalmological Dataset. PLOS Digital Health 3(7):e0000454, 2024.", "https://doi.org/10.1371/journal.pdig.0000454"),
 ("Wu, Restrepo, Nakayama et al. A portable retina fundus photos dataset (mBRSET). Scientific Data 12:340, 2025.", "https://doi.org/10.1038/s41597-025-04627-3"),
]
for i, (txt, url) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    fmt(p.add_run(f"[{i}] {txt} "), 8.5)
    link(p, url)

doc.save(OUT)
print(f"wrote {OUT}")
