"""
Concise technical report (.docx, Times New Roman, black text) for the
ConvNeXt V2 Large multi-label BRSET model (diabetic_retinopathy +
macular_edema). Deliberately short — tables and figures carry the detail,
prose is kept to one or two sentences per section.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Times New Roman"
RESULT_DIR = "/home/users/sthummala2/brset-convnextv2/results/convnextv2_large_BRSET_multilabel_512"
CURVE_FIG = f"{RESULT_DIR}/training_curve.jpg"
CM_FIG = f"{RESULT_DIR}/confusion_matrices.jpg"
OUT_PATH = "/home/users/sthummala2/brset-convnextv2/report/BRSET_ConvNeXtV2_Multilabel_Report.docx"


def set_run_font(run, size=11, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)


def para(doc, text="", size=11, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    sizes = {1: 14, 2: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, size=sizes.get(level, 12), bold=True)
    return p


def make_table(doc, headers, rows, col_widths=None, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size=font_size)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_figure(doc, path, caption, width=5.8):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9.5, italic=True)


def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)

    # ---- Title ----
    para(doc, "ConvNeXt V2 on BRSET: Multi-Label Classification of\nDiabetic Retinopathy and Macular Edema",
         size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "Technical Report", size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc, "Prepared by: Sahith Reddy Thummala  |  Panther ID: 002856791", size=10.5,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Prepared for: Dr. Dong Hye Ye  |  cc: Nagur Shareef Shaik", size=10.5,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Date: July 25, 2026", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # ---- Objective ----
    heading(doc, "1. Objective")
    para(doc, "Per Dr. Ye's direction, this replaces the earlier 5-class ICDR severity task "
              "with independent binary classification of two BRSET findings: diabetic_retinopathy "
              "and macular_edema (each 0/1, an image may carry both, either, or neither).")

    # ---- Dataset ----
    heading(doc, "2. Dataset and Split")
    para(doc, "All 16,266 BRSET images were used, minus 8 confirmed corrupted/missing files "
              "(no quality filter — matches the original paper's own inclusion criteria). "
              "Patient-level 70/15/15 split, stratified jointly on both labels.")
    make_table(doc,
               ["Split", "Patients", "Images", "DR positive", "ME positive"],
               [
                   ["Train", "5,966", "11,372", "748 (6.58%)", "274 (2.41%)"],
                   ["Validation", "1,279", "2,451", "159 (6.49%)", "65 (2.65%)"],
                   ["Test", "1,279", "2,435", "162 (6.65%)", "61 (2.51%)"],
                   ["Total", "8,524", "16,258", "1,069 (6.58%)", "400 (2.46%)"],
               ],
               col_widths=[1.1, 1.1, 1.1, 1.6, 1.6])

    # ---- Method ----
    heading(doc, "3. Model and Training")
    make_table(doc,
               ["Setting", "Value"],
               [
                   ["Backbone", "ConvNeXt V2 Large (196.4M params), convnextv2_large.fcmae_ft_in22k_in1k"],
                   ["Input resolution", "512x512 (fine-tuned above its native 384px pretraining)"],
                   ["Loss", "Multi-label focal loss (gamma=2.0)"],
                   ["Sampling", "Weighted oversampling (inverse label frequency)"],
                   ["Optimizer / LR", "AdamW, lr=5e-5, cosine decay, 3-epoch warmup"],
                   ["Batch size", "16 (x4 gradient accumulation, effective 64)"],
                   ["Epochs planned / used", "40 planned; stopped at epoch 22 (best = epoch 9)"],
                   ["Thresholding", "Per-label threshold tuned on validation set (not 0.5 default)"],
                   ["Inference", "Test-time augmentation (horizontal-flip averaging)"],
                   ["Compute", "1x V100 GPU, SLURM job 3869887, GSU ARC cluster"],
               ],
               col_widths=[2.0, 4.2])
    para(doc, "Training was stopped at epoch 22 rather than continuing to 40: by that point the "
              "learning rate had decayed to ~44% of its peak, training loss had reached 0.0000, "
              "and validation score had not exceeded epoch 9's result in 13 subsequent epochs. "
              "The epoch-9 checkpoint was used for all results below.")
    add_figure(doc, CURVE_FIG, "Figure 1. Validation macro AUC and F1 by epoch. Performance "
                                "plateaus after epoch 9; no further epochs were needed.")

    # ---- Results ----
    heading(doc, "4. Results (Test Set, n=2,435, Epoch-9 Checkpoint, TTA)")
    make_table(doc,
               ["Label", "AUC", "F1", "Precision", "Recall", "Accuracy", "Threshold"],
               [
                   ["diabetic_retinopathy", "0.9872", "0.8452", "0.8161", "0.8765", "0.9786", "0.43"],
                   ["macular_edema", "0.9926", "0.7869", "0.7869", "0.7869", "0.9893", "0.27"],
                   ["Macro average", "0.9899", "0.8161", "-", "-", "-", "-"],
               ],
               col_widths=[2.0, 0.8, 0.8, 0.9, 0.8, 0.9, 0.8])
    add_figure(doc, CM_FIG, "Figure 2. Per-label confusion matrices (counts, row-normalized in "
                             "parentheses), test set.")

    # ---- Paper Comparison ----
    heading(doc, "5. Comparison with the Original BRSET Paper")
    para(doc, "The paper (Nakayama et al., PLOS Digital Health 2024) reports a ConvNeXt V2 "
              "binary DR classifier at AUC 0.97, F1 0.89. It does not report a result for "
              "macular_edema, so that label has no paper baseline to compare against.")
    make_table(doc,
               ["Metric", "Paper (Binary DR)", "This Model (DR)", "Difference"],
               [
                   ["AUC", "0.97", "0.9872", "+0.0172 (better)"],
                   ["F1", "0.89", "0.8452", "-0.0448 (slightly below)"],
               ],
               col_widths=[1.6, 1.8, 1.8, 1.8])

    # ---- Conclusion ----
    heading(doc, "6. Conclusion")
    para(doc, "This is the strongest model produced in this project to date: both labels exceed "
              "AUC 0.98, diabetic_retinopathy AUC surpasses the paper's own published benchmark, "
              "and macular_edema (a rarer, previously unreported finding) reaches AUC 0.99. F1 on "
              "diabetic_retinopathy is close to but below the paper's 0.89, a real gap worth "
              "noting rather than rounding away. Unlike the earlier 5-class ICDR work, no class "
              "collapsed to zero — both labels show real, usable precision and recall.")
    para(doc, "Recommended next step, if pursued further: threshold/recall tuning specifically on "
              "diabetic_retinopathy to close the F1 gap, and re-running the same recipe on the "
              "remaining 11 BRSET classification-parameter labels if broader multi-label coverage "
              "is wanted.")

    doc.save(OUT_PATH)
    print(f"Report written to {OUT_PATH}")


if __name__ == "__main__":
    build()
