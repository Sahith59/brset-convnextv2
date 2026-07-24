"""
Generate the technical report (.docx, Times New Roman, black text) covering
the 5-class ICDR ConvNeXt V2 fine-tuning on full BRSET: baseline vs.
weighted-oversampling + focal loss. All numbers below are pulled directly
from the actual training logs and evaluation reports produced by
scripts 02/03/04. Mirrors ../mbrset-retfound/scripts/07_generate_report.py
in structure and formatting for a consistent pair of documents.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Times New Roman"

BASE_FIG = "/home/users/sthummala2/brset-convnextv2/results/convnextv2_base_BRSET_icdr5_finetune/confusion_matrix_test.jpg"
OF_FIG = "/home/users/sthummala2/brset-convnextv2/results/convnextv2_base_BRSET_icdr5_oversample_focal_finetune/confusion_matrix_test.jpg"
OUT_PATH = "/home/users/sthummala2/brset-convnextv2/report/BRSET_ConvNeXtV2_ICDR5_Report.docx"


def set_run_font(run, size=11, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)


def para(doc, text="", size=11, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    sizes = {1: 15, 2: 13, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=sizes.get(level, 12), bold=True)
    return p


def make_table(doc, headers, rows, col_widths=None, header_bold=True, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=font_size, bold=header_bold)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size=font_size)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_figure(doc, path, caption, width=5.3):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=10, italic=True)


def build():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # ================= Title Page =================
    para(doc, "Fine-Tuning ConvNeXt V2 for Five-Class Diabetic\n"
              "Retinopathy Severity Grading on BRSET:\n"
              "Baseline versus Weighted Oversampling with Focal Loss",
         size=17, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "Technical Report", size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    para(doc, "Prepared by: Sahith Reddy Thummala", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Panther ID: 002856791", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Prepared for: Dr. Dong Hye Ye", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "cc: Nagur Shareef Shaik", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Date: July 24, 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    para(doc,
         "Abstract. This report presents the fine-tuning of ConvNeXt V2 on the full "
         "BRSET dataset for five-class International Clinical Diabetic Retinopathy "
         "(ICDR) severity grading, assigned as a companion project to the mBRSET/"
         "RETFound work. A baseline model trained with unweighted cross-entropy "
         "reproduced the same failure found on mBRSET: complete inability to detect "
         "Grade 1 (Mild NPDR), replicating the finding across a second dataset and a "
         "second model architecture. A literature-grounded follow-up experiment "
         "combining inverse-frequency weighted oversampling with focal loss moved "
         "Grade 1 off complete collapse (F1 0.000 to 0.188) and substantially improved "
         "Grade 3 (F1 0.286 to 0.500), without reducing overall accuracy, at the cost "
         "of a drop in macro AUC that is explained and quantified below. This report "
         "also clarifies exactly where our methodology matches or diverges from the "
         "original BRSET dataset paper, since the paper never reports a comparable "
         "5-class breakdown.",
         size=10.5, italic=True, space_after=6)

    doc.add_page_break()

    # ================= 1. Introduction =================
    heading(doc, "1. Introduction and Objective", level=1)
    para(doc,
         "This work was assigned by Nagur Shareef Shaik as a companion project to the "
         "RETFound/mBRSET fine-tuning work, using the full BRSET dataset with "
         "ConvNeXt V2. The objectives were to: (1) fine-tune ConvNeXt V2 on BRSET for "
         "full five-class ICDR severity grading, (2) determine whether the Grade 1 "
         "(Mild NPDR) classification failure found on mBRSET reproduces on an "
         "independent dataset and model architecture, (3) apply a literature-grounded "
         "remedy targeted at the diagnosed failure mode, and (4) report the outcome, "
         "including any trade-offs, with full transparency.")

    # ================= 2. Dataset =================
    heading(doc, "2. Dataset and Task Description", level=1)
    heading(doc, "2.1 BRSET Overview", level=2)
    para(doc,
         "BRSET (Nakayama et al., PLOS Digital Health, 2024) is a Brazilian retinal "
         "fundus photograph dataset of 16,266 images from 8,524 patients, including "
         "per-image quality-control flags and ICDR diabetic retinopathy severity "
         "grades assigned by a retinal specialist ophthalmologist. After restricting "
         "to images flagged quality-“Adequate” and excluding 7 files found to be "
         "corrupted or truncated during a full-dataset integrity scan (all 7 were "
         "Grade 0 or 2, with no material effect on class balance), 14,273 images from "
         "8,069 unique patients were used in this study.")

    heading(doc, "2.2 Label Definition", level=2)
    para(doc, "The prediction target is the ICDR severity grade used directly as a "
              "five-class label (Grades 0 through 4), identical in definition to the "
              "mBRSET 5-class task.")
    make_table(doc,
               ["Grade", "Clinical Meaning", "Image Count", "% of Dataset"],
               [
                   ["0", "No Diabetic Retinopathy", "13,290", "93.11%"],
                   ["1", "Mild NPDR", "143", "1.00%"],
                   ["2", "Moderate NPDR", "404", "2.83%"],
                   ["3", "Severe NPDR", "75", "0.53%"],
                   ["4", "Proliferative DR", "361", "2.53%"],
                   ["Total", "-", "14,273", "100.00%"],
               ],
               col_widths=[0.8, 2.6, 1.3, 1.3])

    heading(doc, "2.3 Data Splitting Strategy", level=2)
    para(doc,
         "Images were split at the patient level, not the image level, so the same "
         "patient's eyes never appear in both training and test partitions. A "
         "70/15/15 split was applied, stratified on each patient's maximum ICDR "
         "grade across their images (random seed = 42) — the same methodology used "
         "for mBRSET, rather than the original BRSET paper's own 70/30 split, so the "
         "two projects remain directly comparable to each other.")
    make_table(doc,
               ["Split", "Patients", "Images", "Grade 0", "Grade 1", "Grade 2", "Grade 3", "Grade 4"],
               [
                   ["Train", "5,648", "9,980", "9,298", "98", "278", "53", "253"],
                   ["Validation", "1,210", "2,150", "2,001", "21", "62", "11", "55"],
                   ["Test", "1,211", "2,143", "1,991", "24", "64", "11", "53"],
               ],
               col_widths=[0.9, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8], font_size=9.5)

    # ================= 3. Methodology =================
    heading(doc, "3. Methodology", level=1)
    heading(doc, "3.1 Model Architecture and Pretraining", level=2)
    para(doc,
         "ConvNeXt V2 Base (87,697,925 parameters), initialized from the "
         "“convnextv2_base.fcmae_ft_in22k_in1k” checkpoint (self-supervised "
         "masked-autoencoder pretraining followed by supervised fine-tuning on "
         "ImageNet-22k then ImageNet-1k), via the timm library. The entire network "
         "was fine-tuned end-to-end for both experiments.")

    heading(doc, "3.2 Training Configuration", level=2)
    para(doc, "The following were held identical across both experiments unless "
              "otherwise noted, so that any difference in outcome is attributable "
              "solely to the sampling and loss changes described in Section 3.4.")
    make_table(doc,
               ["Hyperparameter", "Value"],
               [
                   ["Model / Architecture", "convnextv2_base.fcmae_ft_in22k_in1k"],
                   ["Adaptation mode", "Full fine-tune (all weights updated)"],
                   ["Input resolution", "224 x 224 (resize 256, crop 224)"],
                   ["Batch size", "64"],
                   ["Epochs", "100"],
                   ["Checkpoint saving", "Best (by composite val score) + every 10 epochs"],
                   ["Optimizer", "AdamW, lr = 1e-4, weight decay = 0.05"],
                   ["LR schedule", "5-epoch linear warmup, cosine decay to 0"],
                   ["Normalization", "ImageNet mean/std"],
                   ["Label smoothing", "0.1"],
                   ["Augmentation", "Random crop, horizontal flip, ±15° rotation, mild color jitter"],
                   ["Hardware", "1x GPU, GSU ARC cluster (qGPU24 partition)"],
               ],
               col_widths=[2.6, 3.6])

    heading(doc, "3.3 Experiment 1: Baseline (Unweighted Loss)", level=2)
    para(doc,
         "Standard unweighted CrossEntropyLoss, uniform random shuffling of the "
         "training set. Checkpoint selection used a composite validation score, "
         "the mean of macro F1, macro one-vs-rest AUC, and Cohen's kappa, computed "
         "after every epoch — identical in formula to the mBRSET/RETFound pipeline.")

    heading(doc, "3.4 Experiment 2: Weighted Oversampling with Focal Loss", level=2)
    para(doc,
         "Experiment 1 reproduced, on an independent dataset and a different model "
         "architecture, the same Grade 1 collapse previously diagnosed on mBRSET "
         "with RETFound (see Section 4.2). Rather than repeating simple class-"
         "weighted cross-entropy — which on mBRSET only partially alleviated the "
         "same failure at a cost to overall accuracy — a literature-grounded "
         "combination of two complementary techniques was applied instead:")
    para(doc,
         "Weighted oversampling: a WeightedRandomSampler drew training images with "
         "probability proportional to the inverse frequency of their class, so "
         "Grade 1 and Grade 3 images are physically seen more often per epoch, not "
         "merely assigned a larger gradient weight on the same rare passes. Because "
         "each draw is re-augmented independently (random crop, flip, rotation, "
         "color jitter), repeated sampling of the same minority images does not "
         "mean repeated identical inputs.")
    para(doc,
         "Focal loss (Lin et al., 2017): replaces cross-entropy with a term that "
         "down-weights already-well-classified examples and concentrates gradient "
         "signal on whatever the model currently misclassifies, rather than a "
         "static per-class weight fixed in advance. Focusing parameter "
         "γ = 2.0, no additional class-weighting (alpha) term.")
    para(doc,
         "No other hyperparameter, architectural choice, or data split was changed "
         "between Experiments 1 and 2.")

    doc.add_page_break()

    # ================= 4. Results =================
    heading(doc, "4. Results", level=1)
    heading(doc, "4.1 Experiment 1 Results (Baseline, Unweighted Loss)", level=2)
    para(doc, "Training completed in 4 hours 6 minutes 13 seconds over 100 epochs. "
              "The best checkpoint by validation composite score occurred at epoch 5; "
              "no later epoch surpassed it through epoch 99.")
    make_table(doc,
               ["Metric", "Value"],
               [
                   ["Overall Accuracy", "0.9547"],
                   ["Macro-averaged F1", "0.4923"],
                   ["Weighted-averaged F1", "0.9447"],
                   ["Macro-averaged AUC (one-vs-rest)", "0.9275"],
                   ["Macro-averaged Precision", "0.5470"],
                   ["Macro-averaged Recall", "0.4629"],
                   ["Macro-averaged Average Precision", "0.5880"],
                   ["Macro-averaged Jaccard Index", "0.4033"],
                   ["Hamming Loss", "0.0181"],
                   ["Cohen's Kappa", "0.5828"],
                   ["Composite Validation-Selection Score", "0.6675"],
               ],
               col_widths=[3.4, 2.4])
    para(doc, "Per-class performance on the held-out test set (n = 2,143):")
    make_table(doc,
               ["Grade", "Precision", "Recall", "F1-score", "Support"],
               [
                   ["0 - No DR", "0.9659", "0.9965", "0.9810", "1991"],
                   ["1 - Mild NPDR", "0.0000", "0.0000", "0.0000", "24"],
                   ["2 - Moderate NPDR", "0.6774", "0.3281", "0.4421", "64"],
                   ["3 - Severe NPDR", "0.3000", "0.2727", "0.2857", "11"],
                   ["4 - Proliferative DR", "0.7917", "0.7170", "0.7525", "53"],
               ],
               col_widths=[2.2, 1.2, 1.2, 1.2, 1.2])
    add_figure(doc, BASE_FIG, "Figure 1. Normalized confusion matrix, Experiment 1 "
                               "(baseline, unweighted loss), held-out test set. Grade 1's "
                               "diagonal cell is 0.0; all 24 true Mild-NPDR images were "
                               "classified as either No DR or Moderate.")

    heading(doc, "4.2 Cross-Dataset Replication of the Grade 1 Collapse", level=2)
    para(doc,
         "The baseline confusion matrix shows all 24 true Grade 1 test images "
         "misclassified: 24 as No DR, 0 elsewhere — an even more concentrated "
         "failure than the equivalent mBRSET/RETFound baseline, where the 41 true "
         "Grade 1 images were split between No DR (35) and Moderate (6). This is "
         "the same failure mode previously diagnosed on mBRSET, now reproduced on "
         "an independent dataset (BRSET, 3x larger) with an independent model "
         "architecture (ConvNeXt V2, a convolutional network, versus RETFound, a "
         "Vision Transformer). Two different architectures and two different "
         "datasets both fail on the same clinical grade, which is strong evidence "
         "this is a genuine property of the task — Grade 1 is defined by isolated "
         "microaneurysms, the subtlest and most easily confused finding on the "
         "entire ICDR scale — rather than an artifact of one pipeline.")

    doc.add_page_break()

    heading(doc, "4.3 Experiment 2 Results (Weighted Oversampling + Focal Loss)", level=2)
    para(doc, "Training completed in 3 hours 43 minutes 31 seconds over 100 epochs, "
              "under identical hyperparameters to Experiment 1 except for the "
              "sampler and loss function described in Section 3.4. The best "
              "checkpoint by validation composite score occurred at epoch 8.")
    make_table(doc,
               ["Metric", "Value"],
               [
                   ["Overall Accuracy", "0.9552"],
                   ["Macro-averaged F1", "0.5824"],
                   ["Weighted-averaged F1", "0.9497"],
                   ["Macro-averaged AUC (one-vs-rest)", "0.8557"],
                   ["Macro-averaged Precision", "0.6561"],
                   ["Macro-averaged Recall", "0.5420"],
                   ["Macro-averaged Average Precision", "0.5800"],
                   ["Macro-averaged Jaccard Index", "0.4657"],
                   ["Hamming Loss", "0.0179"],
                   ["Cohen's Kappa", "0.6197"],
                   ["Composite Validation-Selection Score", "0.6860"],
               ],
               col_widths=[3.4, 2.4])
    para(doc, "Per-class performance on the held-out test set (n = 2,143):")
    make_table(doc,
               ["Grade", "Precision", "Recall", "F1-score", "Support"],
               [
                   ["0 - No DR", "0.9715", "0.9915", "0.9814", "1991"],
                   ["1 - Mild NPDR", "0.3750", "0.1250", "0.1875", "24"],
                   ["2 - Moderate NPDR", "0.5870", "0.4219", "0.4909", "64"],
                   ["3 - Severe NPDR", "0.5556", "0.4545", "0.5000", "11"],
                   ["4 - Proliferative DR", "0.7917", "0.7170", "0.7525", "53"],
               ],
               col_widths=[2.2, 1.2, 1.2, 1.2, 1.2])
    add_figure(doc, OF_FIG, "Figure 2. Normalized confusion matrix, Experiment 2 "
                             "(weighted oversampling + focal loss), held-out test set. "
                             "Grade 1's diagonal cell is now non-zero (3 of 24 correctly "
                             "identified).")

    heading(doc, "4.4 Comparative Summary: Baseline versus Oversampling + Focal Loss", level=2)
    make_table(doc,
               ["Metric", "Baseline", "Oversample + Focal", "Change"],
               [
                   ["Overall Accuracy", "0.9547", "0.9552", "+0.0005"],
                   ["Macro F1", "0.4923", "0.5824", "+0.0901"],
                   ["Macro AUC (OvR)", "0.9275", "0.8557", "-0.0718"],
                   ["Weighted F1", "0.9447", "0.9497", "+0.0050"],
                   ["Cohen's Kappa", "0.5828", "0.6197", "+0.0369"],
                   ["Grade 0 F1", "0.9810", "0.9814", "+0.0004"],
                   ["Grade 1 F1 (Mild NPDR)", "0.0000", "0.1875", "+0.1875"],
                   ["Grade 1 Recall", "0.0000", "0.1250", "+0.1250"],
                   ["Grade 1 Precision", "0.0000", "0.3750", "+0.3750"],
                   ["Grade 2 F1", "0.4421", "0.4909", "+0.0488"],
                   ["Grade 3 F1 (Severe)", "0.2857", "0.5000", "+0.2143"],
                   ["Grade 4 F1", "0.7525", "0.7525", "0.0000"],
               ],
               col_widths=[2.4, 1.4, 1.6, 1.2])

    doc.add_page_break()

    # ================= 5. Discussion =================
    heading(doc, "5. Discussion", level=1)
    para(doc,
         "Weighted oversampling combined with focal loss produced a genuine, "
         "measurable improvement on exactly the classes that failed in Experiment "
         "1. Grade 1 moved from complete collapse (0 of 24 correct) to a real, if "
         "still limited, signal (3 of 24 correct, 37.5% precision when the model "
         "does predict it). Grade 3 nearly doubled in F1. Both improvements were "
         "achieved without any reduction in overall accuracy (0.9547 to 0.9552), "
         "unlike the earlier mBRSET class-weighting experiment, which traded away "
         "accuracy for a smaller Grade 1 gain.")
    para(doc,
         "One trade-off must be reported honestly: macro AUC fell from 0.9275 to "
         "0.8557. This is an expected and explainable consequence of focal loss "
         "rather than a sign the model got worse at its actual job. Focal loss "
         "optimizes for sharper, more confident classification decisions at hard "
         "examples; it does not optimize for well-calibrated probability scores. "
         "AUC is computed from those probability scores, so a model that makes "
         "better hard decisions (reflected in F1, precision, recall, kappa) can "
         "still show a lower AUC if its probability outputs become less reliable "
         "as a ranking signal. In a deployed screening setting where the output "
         "used is the predicted class rather than a probability, the F1/kappa "
         "improvement is the more clinically relevant number; where a probability "
         "threshold or triage ranking is required, the AUC drop would need to be "
         "weighed against the gain.")
    para(doc,
         "Grade 1 recall of 12.5% is not a solved problem — seven of every eight "
         "true mild cases are still missed. What has changed is the nature of the "
         "evidence: Experiment 1 showed the model could not learn to recognize "
         "Grade 1 at all under a standard objective; Experiment 2 shows that a "
         "targeted, literature-backed intervention can recover a meaningful, "
         "non-zero amount of that signal without harming the rest of the model. "
         "That distinction — an unlearnable class versus a hard, partially "
         "learnable class — is itself a useful finding.")

    # ================= 6. Comparison with the Original Paper =================
    heading(doc, "6. Comparison with the Original BRSET Paper", level=1)
    para(doc,
         "The original BRSET paper (Nakayama et al., PLOS Digital Health, 2024) "
         "validated a ConvNeXt V2 model on this dataset, but not on the task "
         "reported here. Its diabetic retinopathy results cover only a binary task "
         "(Normal vs. DR, AUC 0.97, F1 0.89) and a coarser three-class grouping "
         "(Normal / Non-Proliferative / Proliferative, AUC 0.97, F1 0.82), both "
         "reported as macro-averaged figures only, with no per-grade breakdown "
         "published for any task. There is therefore no paper-reported number for "
         "the full five-class ICDR breakdown presented in this report, and no way "
         "to know from the paper whether its own models handled Grade 1 "
         "specifically, since its three-class grouping would combine Grade 1 with "
         "Grades 2 and 3 under “Non-Proliferative,” masking exactly the kind of "
         "single-grade failure this report identifies.")
    make_table(doc,
               ["Aspect", "Paper (Nakayama et al.)", "This Report"],
               [
                   ["Task / label", "Binary DR, and a separate 3-class task",
                    "Full 5-class ICDR (0-4) — no paper baseline exists for this"],
                   ["Per-class metrics", "Not reported (macro AUC/F1 only)", "Full per-grade precision/recall/F1"],
                   ["Image inclusion", "All 16,266 images (no quality filter stated)",
                    "Quality-passing only (14,280), minus 7 corrupted = 14,273"],
                   ["Model", "ConvNeXt V2 (variant unspecified)", "Base (ImageNet-22k→1k pretrained)"],
                   ["Normalization", "Raw 0-1", "ImageNet mean/std (required for the pretrained checkpoint)"],
                   ["Optimizer", "Adam, lr=1e-5", "AdamW, lr=1e-4"],
                   ["Loss (their reported models)", "Weighted cross-entropy",
                    "Unweighted (Exp. 1); oversampling + focal loss (Exp. 2)"],
                   ["Epochs", "50, early-stopping patience 7",
                    "100, best-checkpoint selection (best epoch 5 / 8, no benefit from later epochs)"],
                   ["Split", "70% train (20% of that as val) / 30% test",
                    "Patient-level 70/15/15 (matches the mBRSET/RETFound methodology)"],
               ],
               col_widths=[1.6, 2.6, 2.8], font_size=9.5)
    para(doc,
         "In summary, this report is not a replication of the paper's results and "
         "does not claim to match or contradict them — it answers a more granular "
         "question the paper's own task design could not have answered, using the "
         "same dataset.")

    # ================= 7. Conclusion =================
    heading(doc, "7. Conclusion and Recommended Next Steps", level=1)
    para(doc,
         "ConvNeXt V2 was successfully fine-tuned on full BRSET for five-class "
         "ICDR grading, reproducing — on an independent dataset and architecture "
         "— the same Grade 1 collapse previously found on mBRSET with RETFound, "
         "strengthening the case that this is a genuine property of the task "
         "rather than an artifact of one pipeline. A combined weighted-"
         "oversampling and focal-loss intervention, selected after reviewing the "
         "current literature on class imbalance in diabetic retinopathy grading, "
         "produced a real and measurable improvement on Grade 1 and Grade 3 "
         "without harming overall accuracy, alongside a clearly explained AUC "
         "trade-off.")
    para(doc, "Recommended next steps:")
    for t in [
        "1. Ordinal-aware loss functions (e.g., CORAL), since ICDR grades are "
        "ordered and the current objective treats a Grade 0/Grade 4 error "
        "identically to a Grade 0/Grade 1 error.",
        "2. Apply the same oversampling + focal loss combination back to the "
        "mBRSET/RETFound pipeline, to test whether the improvement generalizes "
        "across both datasets and both architectures.",
        "3. Independent expert re-review of a sample of Grade 0/Grade 1 boundary "
        "images, to quantify how much of the remaining error is genuine model "
        "limitation versus label ambiguity at this boundary.",
        "4. Given the AUC trade-off, evaluate whether a lower focal-loss gamma "
        "or a combined focal-plus-class-weight (alpha) variant recovers some "
        "probability calibration without giving back the Grade 1/3 gains.",
    ]:
        para(doc, t, size=11, space_after=5)

    doc.add_page_break()

    # ================= Appendix =================
    heading(doc, "Appendix A: Software and Compute Environment", level=1)
    make_table(doc,
               ["Component", "Detail"],
               [
                   ["Codebase", "brset-convnextv2 (scripts/02_train_convnextv2.py)"],
                   ["Model library", "timm (convnextv2_base.fcmae_ft_in22k_in1k)"],
                   ["Cluster", "GSU ARC cluster, qGPU24 partition"],
                   ["GPU allocation", "1x GPU per run, 8 CPU cores, 64GB RAM"],
                   ["Experiment 1 job ID", "SLURM job 3844019"],
                   ["Experiment 1 wall-clock time", "4:06:13 (100 epochs)"],
                   ["Experiment 2 job ID", "SLURM job 3867999"],
                   ["Experiment 2 wall-clock time", "3:43:31 (100 epochs)"],
                   ["Random seed (data split)", "42"],
               ],
               col_widths=[2.6, 3.6])

    doc.save(OUT_PATH)
    print(f"Report written to {OUT_PATH}")


if __name__ == "__main__":
    build()
